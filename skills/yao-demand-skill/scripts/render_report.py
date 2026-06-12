#!/usr/bin/env python3
"""Render Yao Demand Skill report JSON to Markdown, HTML, DOCX, and PDF."""

import argparse
import html
import io
import json
import math
import re
import sys
import zipfile
from pathlib import Path
from typing import Any, Dict, Iterable, List, Sequence

from validate_report import validate_report


SECTION_NAV = [
    ("summary", "摘要"),
    ("visuals", "图表"),
    ("product", "产品"),
    ("users", "用户"),
    ("competitors", "竞品"),
    ("triangle", "三角"),
    ("evidence", "证据"),
    ("recommendations", "建议"),
    ("forecast", "预测"),
    ("risks", "风险"),
    ("final", "方案"),
    ("method", "方法"),
    ("appendix", "附录"),
]


def as_list(value: Any) -> List[Any]:
    if value is None:
        return []
    if isinstance(value, list):
        return value
    return [value]


def text(value: Any, default: str = "") -> str:
    if value is None:
        return default
    if isinstance(value, (dict, list)):
        return json.dumps(value, ensure_ascii=False)
    return str(value)


def score(value: Any) -> str:
    try:
        return f"{float(value):.1f}"
    except (TypeError, ValueError):
        return "-"


def slugify(value: str) -> str:
    cleaned = re.sub(r"[^\w\u4e00-\u9fff-]+", "-", value.lower()).strip("-")
    cleaned = re.sub(r"-+", "-", cleaned)
    return cleaned or "demand-report"


def md_escape(value: Any) -> str:
    return text(value).replace("|", "\\|").replace("\n", " ")


def h(value: Any) -> str:
    return html.escape(text(value), quote=True)


def md_list(items: Sequence[Any]) -> str:
    values = [text(item).strip() for item in items if text(item).strip()]
    if not values:
        return "- 未提供\n"
    return "\n".join(f"- {item}" for item in values) + "\n"


def md_table(headers: Sequence[str], rows: Sequence[Sequence[Any]]) -> str:
    if not rows:
        return ""
    header = "| " + " | ".join(md_escape(item) for item in headers) + " |"
    divider = "| " + " | ".join("---" for _ in headers) + " |"
    body = ["| " + " | ".join(md_escape(cell) for cell in row) + " |" for row in rows]
    return "\n".join([header, divider] + body) + "\n"


def html_list(items: Sequence[Any]) -> str:
    values = [h(item) for item in items if text(item).strip()]
    if not values:
        return "<p class=\"muted\">未提供</p>"
    return "<ul>" + "".join(f"<li>{item}</li>" for item in values) + "</ul>"


def html_table(headers: Sequence[str], rows: Sequence[Sequence[Any]], compact: bool = False) -> str:
    if not rows:
        return "<p class=\"muted\">未提供</p>"
    cls = "compact" if compact else ""
    head = "".join(f"<th>{h(item)}</th>" for item in headers)
    body = []
    for row in rows:
        body.append("<tr>" + "".join(f"<td>{h(cell)}</td>" for cell in row) + "</tr>")
    return (
        f"<div class=\"table-wrap\"><table class=\"{cls}\"><thead><tr>{head}</tr></thead>"
        f"<tbody>{''.join(body)}</tbody></table></div>"
    )


def decision_label(value: str) -> str:
    labels = {
        "scale": "可以扩大验证",
        "fix_short_board": "先修短板",
        "validate": "先补验证",
        "reposition": "重做定位",
        "pause": "暂停规模投入",
    }
    return labels.get(value, value or "-")


def score_band(value: Any) -> str:
    try:
        number = float(value)
    except (TypeError, ValueError):
        return "unknown"
    if number >= 8.2:
        return "strong"
    if number >= 6.8:
        return "good"
    if number >= 5.5:
        return "uncertain"
    if number >= 4.0:
        return "fragile"
    return "weak"


def weakest_dimension(report: Dict[str, Any]) -> str:
    summary = report.get("executive_summary", {})
    candidates = [
        ("缺乏感", summary.get("lack_score")),
        ("目标物", summary.get("target_object_score")),
        ("消费者能力", summary.get("consumer_ability_score")),
    ]
    values = [(label, clamp(value)) for label, value in candidates]
    return min(values, key=lambda item: item[1])[0] if values else "-"


def render_process_svg() -> str:
    steps = [
        ("输入", ["链接", "介绍", "文档"]),
        ("解析", ["产品画布", "用户假设", "价值主张"]),
        ("检索", ["权威资料", "竞品", "评价", "价格"]),
        ("分析", ["缺乏感", "目标物", "能力成本"]),
        ("评分", ["三维得分", "信心系数", "红旗项"]),
        ("输出", ["Markdown", "HTML", "Word", "PDF"]),
    ]
    card_w = 130
    card_h = 132
    gap = 36
    parts = [
        '<svg class="process-svg" viewBox="0 0 1030 230" role="img" aria-label="需求评估分析流程图">',
        '<rect x="0" y="0" width="1030" height="230" fill="#ffffff"/>',
    ]
    for index, (title, items) in enumerate(steps):
        x = 26 + index * (card_w + gap)
        parts.append(f'<rect x="{x}" y="32" width="{card_w}" height="{card_h}" rx="6" fill="#ffffff" stroke="#141413" stroke-width="2"/>')
        parts.append(f'<text x="{x + card_w / 2}" y="62" text-anchor="middle" class="svg-title">{h(title)}</text>')
        for item_index, item in enumerate(items):
            parts.append(f'<text x="{x + card_w / 2}" y="{90 + item_index * 18}" text-anchor="middle" class="svg-body">{h(item)}</text>')
        if index < len(steps) - 1:
            ax = x + card_w + 10
            ay = 98
            parts.append(f'<line x1="{ax}" y1="{ay}" x2="{ax + gap - 20}" y2="{ay}" stroke="#141413" stroke-width="2"/>')
            parts.append(f'<path d="M {ax + gap - 20} {ay} l -8 -5 v 10 z" fill="#141413"/>')
    parts.append('<text x="515" y="206" text-anchor="middle" class="svg-caption">核心原则：先把产品放进用户真实场景，再用证据校准三角各边。</text>')
    parts.append("</svg>")
    return "".join(parts)


def render_triangle_svg(report: Dict[str, Any]) -> str:
    summary = report.get("executive_summary", {})
    lack = score(summary.get("lack_score"))
    target = score(summary.get("target_object_score"))
    ability = score(summary.get("consumer_ability_score"))
    total = score(summary.get("total_score"))
    return f"""
<svg class="triangle-svg" viewBox="0 0 760 430" role="img" aria-label="需求三角模型分析结构">
  <rect x="0" y="0" width="760" height="430" fill="#ffffff"/>
  <polygon points="380,82 174,318 586,318" fill="#ffffff" stroke="#1B365D" stroke-width="4" stroke-linejoin="round"/>
  <line x1="380" y1="82" x2="380" y2="318" stroke="#efeee8" stroke-width="1.5"/>
  <line x1="174" y1="318" x2="482" y2="199" stroke="#efeee8" stroke-width="1.5"/>
  <line x1="586" y1="318" x2="278" y2="199" stroke="#efeee8" stroke-width="1.5"/>
  <text x="380" y="34" text-anchor="middle" class="svg-title">缺乏感</text>
  <text x="380" y="57" text-anchor="middle" class="svg-small">理想与现实的差距 | {lack}</text>
  <text x="126" y="358" text-anchor="middle" class="svg-title">目标物</text>
  <text x="126" y="381" text-anchor="middle" class="svg-small">填补差距的方案 | {target}</text>
  <text x="634" y="358" text-anchor="middle" class="svg-title">消费者能力</text>
  <text x="634" y="381" text-anchor="middle" class="svg-small">成本承受力 | {ability}</text>
  <text x="380" y="224" text-anchor="middle" class="svg-title">需求成立</text>
  <text x="380" y="252" text-anchor="middle" class="svg-body">动机清晰 + 成本可承受 + 场景触发</text>
  <text x="380" y="287" text-anchor="middle" class="svg-score">总分 {total}</text>
  <text x="380" y="418" text-anchor="middle" class="svg-caption">任一维度明显缺失，需求成立概率都会下降。</text>
</svg>
"""


def clamp(value: Any, low: float = 0.0, high: float = 10.0) -> float:
    try:
        number = float(value)
    except (TypeError, ValueError):
        return low
    return max(low, min(high, number))


def chart_modules(report: Dict[str, Any]) -> List[Dict[str, Any]]:
    modules = [item for item in as_list(report.get("visual_diagnostics")) if isinstance(item, dict)]
    modules.sort(key=lambda item: (int(item.get("priority", 999)), text(item.get("id"))))
    return modules


def chart_items(module: Dict[str, Any], key: str = "items") -> List[Dict[str, Any]]:
    data = module.get("data", {})
    if not isinstance(data, dict):
        return []
    return [item for item in as_list(data.get(key)) if isinstance(item, dict)]


def chart_meta(module: Dict[str, Any]) -> str:
    source_ids = [text(item) for item in as_list(module.get("source_ids")) if text(item).strip()]
    source_text = " ".join(source_ids) if source_ids else "基于假设"
    return f"置信度 {float(module.get('confidence', 0)):.2f} | {source_text}"


def label_short(value: Any, limit: int = 10) -> str:
    raw = text(value).strip()
    if len(raw) <= limit:
        return raw
    return raw[: max(1, limit - 1)] + "…"


def value_color(value: Any, reverse: bool = False) -> str:
    number = clamp(value)
    if reverse:
        number = 10 - number
    if number >= 7:
        return "#2f6f4e"
    if number >= 5:
        return "#8a641f"
    return "#9b3a32"


def render_score_gauge_svg(module: Dict[str, Any]) -> str:
    data = module.get("data", {}) if isinstance(module.get("data"), dict) else {}
    score_value = clamp(data.get("score", 0))
    decision = h(data.get("decision", ""))
    label = h(data.get("label", "总分"))
    left = 42
    gauge_width = 556
    marker_x = left + (score_value / 10) * gauge_width
    bands = [
        ("弱", 0, 3.9, "#f4e5e2"),
        ("脆弱", 4.0, 5.4, "#f2ead9"),
        ("可验证", 5.5, 6.7, "#eef2f7"),
        ("较强", 6.8, 8.1, "#e7f0ea"),
        ("强", 8.2, 10, "#dfece5"),
    ]
    parts = [
        '<svg viewBox="0 0 640 260" role="img" aria-label="总分诊断图表">',
        '<rect x="0" y="0" width="640" height="260" fill="#ffffff"/>',
        f'<text x="{left}" y="52" font-size="24" fill="#141413" font-family="serif">{label}</text>',
        f'<text x="598" y="52" text-anchor="end" font-size="32" fill="#1B365D" font-family="serif">{score_value:.1f}</text>',
    ]
    for band_label, start, end, color in bands:
        x = left + (start / 10) * gauge_width
        width = ((end - start) / 10) * gauge_width
        parts.append(f'<rect x="{x:.1f}" y="98" width="{width:.1f}" height="40" fill="{color}" stroke="#ffffff"/>')
        parts.append(f'<text x="{x + width / 2:.1f}" y="164" text-anchor="middle" font-size="14" fill="#3d3d3a" font-family="serif">{band_label}</text>')
    parts.append(f'<line x1="{left}" y1="138" x2="{left + gauge_width}" y2="138" stroke="#141413" stroke-width="1"/>')
    parts.append(f'<line x1="{marker_x:.1f}" y1="78" x2="{marker_x:.1f}" y2="150" stroke="#141413" stroke-width="3"/>')
    parts.append(f'<circle cx="{marker_x:.1f}" cy="78" r="7" fill="#1B365D"/>')
    parts.append(f'<text x="{left}" y="214" font-size="18" fill="#141413" font-family="serif">建议动作：{decision}</text>')
    parts.append("</svg>")
    return "".join(parts)


def render_radar_svg(module: Dict[str, Any]) -> str:
    items = chart_items(module)
    if len(items) < 3:
        return render_bar_svg(module)
    cx, cy, radius = 420, 235, 142
    max_value = 10.0
    points = []
    axis_parts = []
    count = len(items)
    for index, item in enumerate(items):
        angle = -math.pi / 2 + index * 2 * math.pi / count
        outer_x = cx + math.cos(angle) * radius
        outer_y = cy + math.sin(angle) * radius
        value = clamp(item.get("value")) / max_value
        point_x = cx + math.cos(angle) * radius * value
        point_y = cy + math.sin(angle) * radius * value
        points.append(f"{point_x:.1f},{point_y:.1f}")
        label_x = cx + math.cos(angle) * (radius + 56)
        label_y = cy + math.sin(angle) * (radius + 42)
        anchor = "middle"
        if label_x < cx - 20:
            anchor = "end"
        elif label_x > cx + 20:
            anchor = "start"
        axis_parts.append(f'<line x1="{cx}" y1="{cy}" x2="{outer_x:.1f}" y2="{outer_y:.1f}" stroke="#e8e6dc" stroke-width="1"/>')
        axis_parts.append(f'<text x="{label_x:.1f}" y="{label_y:.1f}" text-anchor="{anchor}" font-size="15" fill="#3d3d3a" font-family="serif">{h(label_short(item.get("label"), 9))}</text>')
        axis_parts.append(f'<text x="{label_x:.1f}" y="{label_y + 18:.1f}" text-anchor="{anchor}" font-size="13" fill="#1B365D" font-family="serif">{score(item.get("value"))}</text>')
    rings = []
    for factor in [0.25, 0.5, 0.75, 1.0]:
        ring_points = []
        for index in range(count):
            angle = -math.pi / 2 + index * 2 * math.pi / count
            ring_points.append(f"{cx + math.cos(angle) * radius * factor:.1f},{cy + math.sin(angle) * radius * factor:.1f}")
        rings.append(f'<polygon points="{" ".join(ring_points)}" fill="none" stroke="#efeee8" stroke-width="1"/>')
    return (
        '<svg viewBox="0 0 840 470" role="img" aria-label="雷达图">'
        '<rect x="0" y="0" width="840" height="470" fill="#ffffff"/>'
        + "".join(rings)
        + "".join(axis_parts)
        + f'<polygon points="{" ".join(points)}" fill="#EEF2F7" fill-opacity="0.82" stroke="#1B365D" stroke-width="3"/>'
        + f'<circle cx="{cx}" cy="{cy}" r="3" fill="#141413"/>'
        + "</svg>"
    )


def render_bar_svg(module: Dict[str, Any]) -> str:
    items = chart_items(module)
    data = module.get("data", {}) if isinstance(module.get("data"), dict) else {}
    max_value = max(1.0, float(data.get("max", 10) or 10))
    height = max(230, 68 + len(items) * 42)
    left_label_x = 36
    bar_x = 200
    track_width = 346
    value_x = 598
    parts = [
        f'<svg viewBox="0 0 640 {height}" role="img" aria-label="条形图">',
        f'<rect x="0" y="0" width="640" height="{height}" fill="#ffffff"/>',
    ]
    for index, item in enumerate(items):
        y = 40 + index * 42
        value = clamp(item.get("value"), 0, max_value)
        width = (value / max_value) * track_width
        parts.append(f'<text x="{left_label_x}" y="{y + 17}" font-size="15" fill="#141413" font-family="serif">{h(label_short(item.get("label"), 13))}</text>')
        parts.append(f'<rect x="{bar_x}" y="{y}" width="{track_width}" height="20" fill="#f7f7f4" stroke="#efeee8"/>')
        parts.append(f'<rect x="{bar_x}" y="{y}" width="{width:.1f}" height="20" fill="#1B365D"/>')
        parts.append(f'<text x="{value_x}" y="{y + 17}" text-anchor="end" font-size="15" fill="{value_color(value)}" font-family="serif">{value:.1f}</text>')
    parts.append("</svg>")
    return "".join(parts)


def render_heatmap_svg(module: Dict[str, Any]) -> str:
    items = chart_items(module)
    groups: Dict[str, List[Dict[str, Any]]] = {}
    for item in items:
        groups.setdefault(text(item.get("group", "其他")), []).append(item)
    cell_w, cell_h = 96, 56
    left, top = 132, 42
    max_cols = max((len(values) for values in groups.values()), default=1)
    height = top + len(groups) * (cell_h + 26) + 30
    width = max(840, left + max_cols * cell_w + 50)
    parts = [
        f'<svg viewBox="0 0 {width} {height}" role="img" aria-label="热力图">',
        f'<rect x="0" y="0" width="{width}" height="{height}" fill="#ffffff"/>',
    ]
    for row_index, (group, values) in enumerate(groups.items()):
        y = top + row_index * (cell_h + 26)
        parts.append(f'<text x="36" y="{y + 32}" font-size="16" fill="#141413" font-family="serif">{h(group)}</text>')
        for col_index, item in enumerate(values):
            x = left + col_index * cell_w
            value = clamp(item.get("value"))
            fill = "#e7f0ea" if value >= 7 else "#f2ead9" if value >= 5 else "#f4e5e2"
            parts.append(f'<rect x="{x}" y="{y}" width="{cell_w - 8}" height="{cell_h}" rx="5" fill="{fill}" stroke="#e8e6dc"/>')
            parts.append(f'<text x="{x + (cell_w - 8) / 2}" y="{y + 22}" text-anchor="middle" font-size="13" fill="#3d3d3a" font-family="serif">{h(label_short(item.get("label"), 6))}</text>')
            parts.append(f'<text x="{x + (cell_w - 8) / 2}" y="{y + 44}" text-anchor="middle" font-size="18" fill="{value_color(value)}" font-family="serif">{value:.0f}</text>')
    parts.append("</svg>")
    return "".join(parts)


def render_matrix_svg(module: Dict[str, Any]) -> str:
    data = module.get("data", {}) if isinstance(module.get("data"), dict) else {}
    items = chart_items(module)
    x_axis = h(data.get("x_axis", "横轴"))
    y_axis = h(data.get("y_axis", "纵轴"))
    left, top, width, height = 74, 44, 500, 250
    points = []
    for index, item in enumerate(items):
        x = left + (clamp(item.get("x")) / 10) * width
        y = top + height - (clamp(item.get("y")) / 10) * height
        r = 7.0 + clamp(item.get("size", 5), 0, 10) * 0.35
        points.append(
            {
                "item": item,
                "x": x,
                "y": y,
                "r": r,
                "index": index + 1,
            }
        )
    min_dist = 22
    for _ in range(6):
        for i in range(len(points)):
            for j in range(i + 1, len(points)):
                dx = points[j]["x"] - points[i]["x"]
                dy = points[j]["y"] - points[i]["y"]
                dist = math.hypot(dx, dy)
                if dist >= min_dist:
                    continue
                if dist < 0.1:
                    angle = (i + j + 1) * 0.9
                    dx = math.cos(angle)
                    dy = math.sin(angle)
                    dist = 1.0
                push = (min_dist - dist) / 2
                ux, uy = dx / dist, dy / dist
                points[i]["x"] -= ux * push
                points[i]["y"] -= uy * push
                points[j]["x"] += ux * push
                points[j]["y"] += uy * push
                for point in (points[i], points[j]):
                    r = point["r"] + 1
                    point["x"] = max(left + r, min(left + width - r, point["x"]))
                    point["y"] = max(top + r, min(top + height - r, point["y"]))
    legend_top = 356
    legend_row_h = 28
    legend_cols = 2
    legend_height = max(1, math.ceil(len(points) / legend_cols)) * legend_row_h + 24
    svg_height = max(470, legend_top + legend_height)
    parts = [
        f'<svg viewBox="0 0 640 {svg_height}" role="img" aria-label="矩阵图">',
        f'<rect x="0" y="0" width="640" height="{svg_height}" fill="#ffffff"/>',
        f'<rect x="{left}" y="{top}" width="{width}" height="{height}" fill="#ffffff" stroke="#141413" stroke-width="1.5"/>',
        f'<line x1="{left + width / 2}" y1="{top}" x2="{left + width / 2}" y2="{top + height}" stroke="#efeee8" stroke-width="2"/>',
        f'<line x1="{left}" y1="{top + height / 2}" x2="{left + width}" y2="{top + height / 2}" stroke="#efeee8" stroke-width="2"/>',
        f'<text x="{left + width / 2}" y="328" text-anchor="middle" font-size="15" fill="#3d3d3a" font-family="serif">{x_axis}</text>',
        f'<text x="22" y="{top + height / 2}" transform="rotate(-90 22 {top + height / 2})" text-anchor="middle" font-size="15" fill="#3d3d3a" font-family="serif">{y_axis}</text>',
        f'<text x="{left + 8}" y="{top + 20}" font-size="12" fill="#8a641f" font-family="serif">低确定性</text>',
        f'<text x="{left + width - 8}" y="{top + 20}" text-anchor="end" font-size="12" fill="#2f6f4e" font-family="serif">高确定性</text>',
        f'<line x1="{left}" y1="{legend_top - 18}" x2="{left + width}" y2="{legend_top - 18}" stroke="#efeee8" stroke-width="1"/>',
    ]
    for point in points:
        x = point["x"]
        y = point["y"]
        r = point["r"]
        marker = point["index"]
        parts.append(f'<circle cx="{x:.1f}" cy="{y:.1f}" r="{r:.1f}" fill="#ffffff" stroke="#1B365D" stroke-width="2"/>')
        parts.append(
            f'<text x="{x:.1f}" y="{y + 3.6:.1f}" text-anchor="middle" font-size="10" '
            f'fill="#1B365D" font-family="serif">{marker}</text>'
        )
    for point in points:
        item = point["item"]
        index = point["index"]
        col = (index - 1) % legend_cols
        row = (index - 1) // legend_cols
        legend_x = left + col * 270
        legend_y = legend_top + row * legend_row_h
        parts.append(f'<circle cx="{legend_x}" cy="{legend_y - 4}" r="8" fill="#ffffff" stroke="#1B365D" stroke-width="1.5"/>')
        parts.append(
            f'<text x="{legend_x}" y="{legend_y}" text-anchor="middle" font-size="9" fill="#1B365D" font-family="serif">{index}</text>'
        )
        parts.append(
            f'<text x="{legend_x + 16}" y="{legend_y}" font-size="12" fill="#141413" font-family="serif">{h(label_short(item.get("label"), 12))}</text>'
        )
        parts.append(
            f'<text x="{legend_x + 172}" y="{legend_y}" font-size="10.5" fill="#6b6a64" font-family="serif">{score(item.get("x"))} / {score(item.get("y"))}</text>'
        )
    parts.append("</svg>")
    return "".join(parts)


def render_funnel_svg(module: Dict[str, Any]) -> str:
    items = chart_items(module)
    max_value = max([clamp(item.get("value"), 0, 100) for item in items] or [100])
    height = max(250, 54 + len(items) * 48)
    parts = [
        f'<svg viewBox="0 0 840 {height}" role="img" aria-label="漏斗图">',
        f'<rect x="0" y="0" width="840" height="{height}" fill="#ffffff"/>',
    ]
    for index, item in enumerate(items):
        value = clamp(item.get("value"), 0, max_value)
        bar_w = 560 * (value / max_value)
        x = 250 + (560 - bar_w) / 2
        y = 34 + index * 46
        parts.append(f'<text x="42" y="{y + 25}" font-size="15" fill="#141413" font-family="serif">{h(label_short(item.get("label"), 14))}</text>')
        parts.append(f'<rect x="{x:.1f}" y="{y}" width="{bar_w:.1f}" height="28" rx="4" fill="#1B365D" fill-opacity="{0.95 - index * 0.055:.2f}"/>')
        parts.append(f'<text x="770" y="{y + 21}" text-anchor="end" font-size="15" fill="#3d3d3a" font-family="serif">{value:.0f}</text>')
    parts.append("</svg>")
    return "".join(parts)


def render_stacked_bar_svg(module: Dict[str, Any]) -> str:
    data = module.get("data", {}) if isinstance(module.get("data"), dict) else {}
    segments = [item for item in as_list(data.get("segments")) if isinstance(item, dict)]
    total = sum(max(0, float(item.get("value", 0) or 0)) for item in segments) or 1
    colors = ["#1B365D", "#7A8CA5", "#C7D0DB", "#e8e6dc", "#8a641f"]
    x = 70
    parts = [
        '<svg viewBox="0 0 840 260" role="img" aria-label="堆叠条形图">',
        '<rect x="0" y="0" width="840" height="260" fill="#ffffff"/>',
        '<rect x="70" y="84" width="700" height="44" fill="#f7f7f4" stroke="#efeee8"/>',
    ]
    for index, item in enumerate(segments):
        value = max(0, float(item.get("value", 0) or 0))
        width = 700 * value / total
        parts.append(f'<rect x="{x:.1f}" y="84" width="{width:.1f}" height="44" fill="{colors[index % len(colors)]}"/>')
        if width > 70:
            parts.append(f'<text x="{x + width / 2:.1f}" y="112" text-anchor="middle" font-size="14" fill="#ffffff" font-family="serif">{h(label_short(item.get("label"), 8))}</text>')
        x += width
    legend_x = 70
    for index, item in enumerate(segments):
        y = 164 + (index // 3) * 32
        lx = legend_x + (index % 3) * 230
        parts.append(f'<rect x="{lx}" y="{y}" width="12" height="12" fill="{colors[index % len(colors)]}"/>')
        parts.append(f'<text x="{lx + 18}" y="{y + 12}" font-size="14" fill="#141413" font-family="serif">{h(label_short(item.get("label"), 12))} {text(item.get("value"))}</text>')
    parts.append("</svg>")
    return "".join(parts)


def render_forecast_svg(module: Dict[str, Any]) -> str:
    data = module.get("data", {}) if isinstance(module.get("data"), dict) else {}
    scenarios = [item for item in as_list(data.get("scenarios")) if isinstance(item, dict)]
    if not scenarios:
        scenarios = chart_items(module)
    left, top, width, height = 90, 48, 650, 280
    parts = [
        '<svg viewBox="0 0 840 430" role="img" aria-label="预测情景图">',
        '<rect x="0" y="0" width="840" height="430" fill="#ffffff"/>',
        f'<line x1="{left}" y1="{top + height}" x2="{left + width}" y2="{top + height}" stroke="#141413" stroke-width="1.5"/>',
        f'<line x1="{left}" y1="{top}" x2="{left}" y2="{top + height}" stroke="#141413" stroke-width="1.5"/>',
    ]
    points = []
    count = max(1, len(scenarios) - 1)
    for index, scenario in enumerate(scenarios):
        x = left + (index / count) * width if count else left
        y = top + height - (clamp(scenario.get("score_after")) / 10) * height
        points.append(f"{x:.1f},{y:.1f}")
        parts.append(f'<circle cx="{x:.1f}" cy="{y:.1f}" r="8" fill="#1B365D"/>')
        parts.append(f'<text x="{x:.1f}" y="{y - 14:.1f}" text-anchor="middle" font-size="15" fill="#141413" font-family="serif">{score(scenario.get("score_after"))}</text>')
        parts.append(f'<text x="{x:.1f}" y="370" text-anchor="middle" font-size="13" fill="#3d3d3a" font-family="serif">{h(label_short(scenario.get("name"), 10))}</text>')
        parts.append(f'<text x="{x:.1f}" y="390" text-anchor="middle" font-size="12" fill="#6b6a64" font-family="serif">{h(scenario.get("adoption_likelihood", ""))}</text>')
    if len(points) > 1:
        parts.append(f'<polyline points="{" ".join(points)}" fill="none" stroke="#1B365D" stroke-width="3"/>')
    parts.append('<text x="40" y="52" font-size="13" fill="#6b6a64" font-family="serif">10</text>')
    parts.append('<text x="48" y="330" font-size="13" fill="#6b6a64" font-family="serif">0</text>')
    parts.append("</svg>")
    return "".join(parts)


def render_chart_svg(module: Dict[str, Any]) -> str:
    chart_type = module.get("chart_type")
    if chart_type == "score_gauge":
        return render_score_gauge_svg(module)
    if chart_type == "radar":
        return render_radar_svg(module)
    if chart_type == "bar":
        return render_bar_svg(module)
    if chart_type == "heatmap":
        return render_heatmap_svg(module)
    if chart_type == "matrix":
        return render_matrix_svg(module)
    if chart_type == "funnel":
        return render_funnel_svg(module)
    if chart_type == "stacked_bar":
        return render_stacked_bar_svg(module)
    if chart_type == "forecast":
        return render_forecast_svg(module)
    return render_bar_svg(module)


def chart_data_rows(module: Dict[str, Any]) -> List[List[Any]]:
    data = module.get("data", {}) if isinstance(module.get("data"), dict) else {}
    if "items" in data:
        rows = []
        for item in chart_items(module):
            if "value" in item:
                rows.append([item.get("group", ""), item.get("label", ""), score(item.get("value")), ""])
            else:
                rows.append(["", item.get("label", ""), score(item.get("x")), score(item.get("y"))])
        return rows
    if "segments" in data:
        return [["", item.get("label", ""), item.get("value", ""), ""] for item in as_list(data.get("segments")) if isinstance(item, dict)]
    if "scenarios" in data:
        return [["", item.get("name", ""), score(item.get("score_after")), item.get("adoption_likelihood", "")] for item in as_list(data.get("scenarios")) if isinstance(item, dict)]
    if "score" in data:
        return [["", data.get("label", "分数"), score(data.get("score")), data.get("decision", "")]]
    return [["", "数据", text(data)]]


def render_chart_module_html(module: Dict[str, Any]) -> str:
    chart_class = f"chart-{h(module.get('chart_type'))}"
    return f"""
    <article class="chart-module {chart_class}" id="chart-{h(module.get('id'))}">
      <div class="chart-module-head">
        <div>
          <span class="chart-kicker">{h(module.get('chart_type'))}</span>
          <h3>{h(module.get('title'))}</h3>
        </div>
        <span class="chart-confidence">{h(chart_meta(module))}</span>
      </div>
      <div class="chart-svg-wrap">{render_chart_svg(module)}</div>
      <p class="chart-insight"><strong>解读：</strong>{h(module.get('insight'))}</p>
      <p class="chart-recommendation"><strong>建议：</strong>{h(module.get('recommendation'))}</p>
    </article>
    """


def score_rows(report: Dict[str, Any]) -> List[List[Any]]:
    summary = report.get("executive_summary", {})
    return [
        ["缺乏感", score(summary.get("lack_score"))],
        ["目标物", score(summary.get("target_object_score"))],
        ["消费者能力", score(summary.get("consumer_ability_score"))],
        ["证据信心", text(summary.get("evidence_confidence", "-"))],
        ["总分", score(summary.get("total_score"))],
    ]


def competitor_rows(report: Dict[str, Any]) -> List[List[Any]]:
    rows = []
    for item in as_list(report.get("competitors")):
        rows.append(
            [
                item.get("name", ""),
                item.get("type", ""),
                item.get("positioning", ""),
                "；".join(as_list(item.get("strengths"))),
                "；".join(as_list(item.get("weaknesses"))),
                " ".join(as_list(item.get("source_ids"))),
            ]
        )
    return rows


def segment_rows(report: Dict[str, Any]) -> List[List[Any]]:
    rows = []
    for item in as_list(report.get("segments")):
        rows.append(
            [
                item.get("name", ""),
                item.get("scenario", ""),
                item.get("job_to_be_done", ""),
                "；".join(as_list(item.get("current_alternatives"))),
                "；".join(as_list(item.get("adoption_blockers"))),
            ]
        )
    return rows


def evidence_rows(report: Dict[str, Any]) -> List[List[Any]]:
    rows = []
    for item in as_list(report.get("evidence")):
        rows.append(
            [
                item.get("id", ""),
                item.get("quality", ""),
                item.get("source_type", ""),
                item.get("title", ""),
                item.get("url", ""),
                item.get("retrieved_at") or item.get("published_at", ""),
            ]
        )
    return rows


def risk_rows(report: Dict[str, Any]) -> List[List[Any]]:
    rows = []
    for item in as_list(report.get("risks")):
        rows.append(
            [
                item.get("severity", ""),
                item.get("type", ""),
                item.get("risk", ""),
                item.get("mitigation", ""),
                " ".join(as_list(item.get("source_ids"))),
            ]
        )
    return rows


def experiment_rows(report: Dict[str, Any]) -> List[List[Any]]:
    rows = []
    for item in as_list(report.get("experiments")):
        rows.append(
            [
                item.get("hypothesis", ""),
                item.get("segment", ""),
                item.get("method", ""),
                item.get("metric", ""),
                item.get("threshold", ""),
                item.get("decision_rule", ""),
            ]
        )
    return rows


def render_markdown(report: Dict[str, Any]) -> str:
    meta = report.get("meta", {})
    summary = report.get("executive_summary", {})
    canvas = report.get("product_canvas", {})
    triangle = report.get("triangle_analysis", {})
    lines: List[str] = []

    lines.append(f"# {text(meta.get('title') or meta.get('product_name') or '需求评估报告')}\n")
    lines.append(f"- 产品：{text(meta.get('product_name'))}")
    lines.append(f"- 日期：{text(meta.get('generated_at'))}")
    lines.append(f"- 目标市场：{text(meta.get('target_market', '未指定'))}")
    lines.append(f"- 分析目标：{text(meta.get('analysis_goal', '需求评估'))}")
    lines.append(f"- 来源边界：{text(meta.get('source_boundary', '未说明'))}\n")

    lines.append("## 执行摘要\n")
    lines.append(f"{text(summary.get('one_sentence'))}\n")
    lines.append(md_table(["项目", "结果"], [
        ["建议决策", decision_label(summary.get("decision", ""))],
        ["总分", score(summary.get("total_score"))],
        ["缺乏感", score(summary.get("lack_score"))],
        ["目标物", score(summary.get("target_object_score"))],
        ["消费者能力", score(summary.get("consumer_ability_score"))],
        ["证据信心", text(summary.get("evidence_confidence"))],
        ["最大机会", summary.get("biggest_opportunity", "")],
        ["最大风险", summary.get("biggest_risk", "")],
    ]))

    lines.append("## 可视化诊断\n")
    lines.append("以下图表模块用于快速定位需求强弱、短板、证据质量、采用阻力和下一步优先级。Markdown 版本提供图表等价表格，HTML/PDF 会渲染为静态 SVG 图表。\n")
    for module in chart_modules(report):
        lines.append(f"### {text(module.get('title'))}\n")
        lines.append(f"- 图表类型：`{text(module.get('chart_type'))}`")
        lines.append(f"- {chart_meta(module)}")
        lines.append(f"- 解读：{text(module.get('insight'))}")
        lines.append(f"- 建议：{text(module.get('recommendation'))}\n")
        lines.append(md_table(["分组", "指标/情景", "数值/X", "Y/说明"], chart_data_rows(module)))

    lines.append("## 产品概览\n")
    lines.append(f"**产品定义：** {text(canvas.get('definition'))}\n")
    lines.append(f"**价值主张：** {text(canvas.get('value_proposition'))}\n")
    lines.append("**核心功能：**\n")
    lines.append(md_list(as_list(canvas.get("features"))))
    lines.append("**定价与商业模式：**\n")
    lines.append(md_list(as_list(canvas.get("pricing")) + [canvas.get("business_model", "")]))
    lines.append("**假设：**\n")
    lines.append(md_list(as_list(canvas.get("assumptions"))))

    lines.append("## 研究方法与来源\n")
    lines.append(text(meta.get("source_boundary", "本报告区分事实、假设、证据和建议。")) + "\n")
    lines.append("```mermaid\nflowchart LR\n  A[输入\\n链接/介绍/文档] --> B[解析\\n产品画布/用户假设/价值主张]\n  B --> C[检索\\n权威资料/竞品/评价/价格]\n  C --> D[分析\\n缺乏感/目标物/能力成本]\n  D --> E[评分\\n三维得分/信心系数/红旗项]\n  E --> F[输出\\nMarkdown/HTML/Word/PDF]\n```\n")
    lines.append(md_table(["来源", "等级", "类型", "标题", "链接", "日期"], evidence_rows(report)))

    lines.append("## 目标用户与 JTBD\n")
    lines.append(md_table(["分群", "场景", "JTBD", "当前替代", "采用阻碍"], segment_rows(report)))

    lines.append("## 竞品与替代方案\n")
    lines.append(md_table(["名称", "类型", "定位", "优势", "弱点", "来源"], competitor_rows(report)))

    lines.append("## 需求三角分析\n")
    lines.append("```mermaid\nflowchart TB\n  L[缺乏感\\n理想与现实的差距] --> M[需求成立\\n动机清晰 + 成本可承受 + 场景触发]\n  T[目标物\\n填补差距的具体方案] --> M\n  A[消费者能力\\n行动所需成本承受力] --> M\n```\n")
    for key, label in [("lack", "缺乏感"), ("target_object", "目标物"), ("consumer_ability", "消费者能力")]:
        dim = triangle.get(key, {})
        lines.append(f"### {label}：{score(dim.get('score'))}\n")
        lines.append(f"**推理：** {text(dim.get('reasoning'))}\n")
        lines.append("**支持证据：**\n")
        lines.append(md_list(as_list(dim.get("evidence"))))
        lines.append("**反证或缺口：**\n")
        lines.append(md_list(as_list(dim.get("counter_evidence"))))
        lines.append(f"**改进路径：** {text(dim.get('improvement_path'))}\n")

    lines.append("## 评分与解释\n")
    lines.append(md_table(["维度", "分数"], score_rows(report)))

    lines.append("## 建议与实验\n")
    rec_rows = [
        [
            item.get("priority", ""),
            item.get("area", ""),
            item.get("recommendation", ""),
            item.get("rationale", ""),
            item.get("expected_impact", ""),
            item.get("effort", ""),
        ]
        for item in as_list(report.get("recommendations"))
    ]
    lines.append(md_table(["优先级", "领域", "建议", "理由", "预期影响", "成本"], rec_rows))
    lines.append(md_table(["假设", "分群", "方法", "指标", "阈值", "决策规则"], experiment_rows(report)))

    lines.append("## 风险与伦理\n")
    lines.append(md_table(["严重度", "类型", "风险", "缓释", "来源"], risk_rows(report)))

    forecast = report.get("forecast", {})
    lines.append("## 预测情景\n")
    lines.append(f"- 预测窗口：{text(forecast.get('horizon', '未提供'))}")
    lines.append(f"- 置信度：{text(forecast.get('confidence', '未提供'))}")
    lines.append(f"- 复盘触发：{text(forecast.get('recheck_trigger', '未提供'))}\n")
    lines.append(md_table(
        ["情景", "预测分数", "采用可能性", "关键假设"],
        [
            [
                item.get("name", ""),
                score(item.get("score_after")),
                item.get("adoption_likelihood", ""),
                "；".join(as_list(item.get("assumptions"))),
            ]
            for item in as_list(forecast.get("scenarios"))
            if isinstance(item, dict)
        ],
    ))

    final_plan = report.get("final_plan", {})
    lines.append("## 最终方案\n")
    lines.append(f"**最终判断：** {text(final_plan.get('final_judgment'))}\n")
    lines.append(f"**总体策略：** {text(final_plan.get('strategy'))}\n")
    for key, label in [("next_30_days", "未来 30 天"), ("next_60_days", "未来 60 天"), ("next_90_days", "未来 90 天"), ("decision_rules", "决策规则")]:
        lines.append(f"### {label}\n")
        lines.append(md_list(as_list(final_plan.get(key))))

    lines.append("## 附录\n")
    lines.append("### 未解问题\n")
    lines.append(md_list(as_list(report.get("open_questions"))))
    return "\n".join(lines).rstrip() + "\n"


def render_html(report: Dict[str, Any]) -> str:
    meta = report.get("meta", {})
    summary = report.get("executive_summary", {})
    canvas = report.get("product_canvas", {})
    triangle = report.get("triangle_analysis", {})
    visual_modules = chart_modules(report)
    forecast = report.get("forecast", {})
    final_plan = report.get("final_plan", {})
    title = meta.get("title") or meta.get("product_name") or "需求评估报告"
    band = score_band(summary.get("total_score"))
    nav_links = "".join(f"<a href=\"#{anchor}\">{label}</a>" for anchor, label in SECTION_NAV)
    chart_html = "\n".join(render_chart_module_html(module) for module in visual_modules)

    dimension_cards = []
    for key, label in [("lack", "缺乏感"), ("target_object", "目标物"), ("consumer_ability", "消费者能力")]:
        dim = triangle.get(key, {})
        dimension_cards.append(
            f"""
            <article class="dimension-card">
              <div class="dimension-head"><h3>{h(label)}</h3><span>{score(dim.get('score'))}</span></div>
              <p>{h(dim.get('reasoning'))}</p>
              <h4>支持证据</h4>
              {html_list(as_list(dim.get('evidence')))}
              <h4>反证或缺口</h4>
              {html_list(as_list(dim.get('counter_evidence')))}
              <div class="note"><strong>改进路径</strong><br>{h(dim.get('improvement_path'))}</div>
            </article>
            """
        )

    rec_rows = [
        [
            item.get("priority", ""),
            item.get("area", ""),
            item.get("recommendation", ""),
            item.get("rationale", ""),
            item.get("expected_impact", ""),
            item.get("effort", ""),
        ]
        for item in as_list(report.get("recommendations"))
    ]

    html_doc = f"""<!doctype html>
<html lang="{h(meta.get('language', 'zh-CN'))}">
<head>
  <meta charset="utf-8">
  <meta name="viewport" content="width=device-width, initial-scale=1">
  <title>{h(title)}</title>
  <meta name="generator" content="Yao Demand Skill">
  <style>
    :root {{
      --paper: #ffffff;
      --near-black: #141413;
      --dark-warm: #3d3d3a;
      --olive: #504e49;
      --stone: #6b6a64;
      --border: #e8e6dc;
      --border-soft: #efeee8;
      --brand: #1B365D;
      --brand-soft: #EEF2F7;
      --green: #2f6f4e;
      --amber: #8a641f;
      --red: #9b3a32;
      --serif: "TsangerJinKai02", "Source Han Serif SC", "Noto Serif CJK SC", "Songti SC", "STSong", Georgia, serif;
      --mono: "JetBrains Mono", "SF Mono", Consolas, "TsangerJinKai02", "Source Han Serif SC", monospace;
    }}
    * {{ box-sizing: border-box; }}
    html {{ scroll-behavior: smooth; scroll-padding-top: 78px; background: var(--paper); overflow-x: hidden; }}
    body {{
      margin: 0;
      padding-top: 57px;
      background: var(--paper);
      color: var(--near-black);
      font-family: var(--serif);
      font-size: 16px;
      line-height: 1.62;
      letter-spacing: 0;
      overflow-x: hidden;
      text-rendering: optimizeLegibility;
    }}
    .top-nav {{
      position: fixed;
      top: 0;
      left: 0;
      right: 0;
      z-index: 100;
      background: #ffffff;
      border-bottom: 1px solid var(--border);
      box-shadow: 0 2px 14px rgba(20, 20, 19, 0.06);
    }}
    .nav-inner {{
      max-width: 1120px;
      margin: 0 auto;
      padding: 10px 28px;
      display: flex;
      align-items: center;
      gap: 24px;
    }}
    .nav-title {{
      min-width: 180px;
      max-width: 320px;
      color: var(--near-black);
      font-size: 15px;
      line-height: 1.25;
      font-weight: 500;
      white-space: nowrap;
      overflow: hidden;
      text-overflow: ellipsis;
    }}
    .nav-links {{
      display: flex;
      gap: 14px;
      overflow-x: auto;
      scrollbar-width: thin;
      white-space: nowrap;
    }}
    .nav-links a {{
      color: var(--brand);
      text-decoration: none;
      font-size: 14px;
      padding: 5px 0 7px;
      border-bottom: 2px solid transparent;
    }}
    .nav-links a:hover {{ border-bottom-color: var(--brand); }}
    main {{
      max-width: 1180px;
      margin: 0 auto;
      padding: 48px 28px 72px;
    }}
    header.report-header {{
      border-bottom: 1px solid var(--border);
      padding-bottom: 28px;
      margin-bottom: 28px;
    }}
    .eyebrow {{
      color: var(--brand);
      font-size: 13px;
      letter-spacing: 0.5px;
      margin-bottom: 12px;
    }}
    h1 {{
      margin: 0 0 14px;
      font-size: clamp(32px, 5vw, 54px);
      line-height: 1.12;
      font-weight: 500;
    }}
    h2 {{
      margin: 0 0 16px;
      font-size: 26px;
      line-height: 1.25;
      font-weight: 500;
      border-left: 4px solid var(--brand);
      padding-left: 12px;
    }}
    h3 {{
      margin: 0 0 8px;
      font-size: 19px;
      line-height: 1.3;
      font-weight: 500;
    }}
    h4 {{
      margin: 18px 0 8px;
      font-size: 15px;
      line-height: 1.35;
      color: var(--dark-warm);
      font-weight: 500;
    }}
    section {{
      padding: 46px 0;
      border-top: 1px solid var(--border-soft);
      scroll-margin-top: 74px;
    }}
    section + section {{
      margin-top: 8px;
    }}
    .lede {{
      max-width: 980px;
      color: var(--dark-warm);
      font-size: 18px;
      line-height: 1.72;
      text-wrap: pretty;
    }}
    .meta-grid, .score-grid, .fact-grid {{
      display: grid;
      grid-template-columns: repeat(4, minmax(0, 1fr));
      gap: 14px;
      margin-top: 22px;
      align-items: start;
    }}
    .meta-item, .score-card, .fact-box, .dimension-card, .note {{
      border: 1px solid var(--border);
      background: var(--paper);
      border-radius: 0;
      padding: 16px;
    }}
    .meta-item span, .score-card span {{
      display: block;
      color: var(--stone);
      font-size: 13px;
      margin-bottom: 6px;
    }}
    .score-card strong {{
      display: block;
      font-size: 28px;
      line-height: 1.1;
      font-weight: 500;
      color: var(--brand);
    }}
    .decision {{
      display: inline-block;
      color: var(--brand);
      background: var(--brand-soft);
      border-radius: 0;
      padding: 2px 8px;
      font-size: 14px;
    }}
    .total-score {{
      border-top: 4px solid var(--brand);
    }}
    .band-strong strong, .band-good strong {{ color: var(--green); }}
    .band-uncertain strong {{ color: var(--amber); }}
    .band-fragile strong, .band-weak strong {{ color: var(--red); }}
    .dimension-grid {{
      display: grid;
      grid-template-columns: repeat(3, minmax(0, 1fr));
      gap: 18px;
    }}
    .diagram-block {{
      margin: 18px 0 26px;
      padding: 0;
      overflow-x: auto;
    }}
    .diagram-block svg {{
      width: 100%;
      min-width: 720px;
      height: auto;
      display: block;
    }}
    .triangle-figure {{
      text-align: center;
    }}
    .triangle-figure svg {{
      max-width: 760px;
      min-width: 0;
      margin: 0 auto;
    }}
    .figure-caption {{
      margin-top: 8px;
      color: var(--stone);
      font-size: 14px;
      line-height: 1.45;
    }}
    .svg-title {{
      font-family: var(--serif);
      font-size: 19px;
      font-weight: 500;
      fill: #141413;
    }}
    .svg-body {{
      font-family: var(--serif);
      font-size: 17px;
      fill: #141413;
    }}
    .svg-small {{
      font-family: var(--serif);
      font-size: 13px;
      fill: #3d3d3a;
    }}
    .svg-score {{
      font-family: var(--serif);
      font-size: 22px;
      font-weight: 500;
      fill: #1B365D;
    }}
    .svg-caption {{
      font-family: var(--serif);
      font-size: 15px;
      font-weight: 500;
      fill: #141413;
    }}
    .dimension-head {{
      display: flex;
      justify-content: space-between;
      align-items: baseline;
      gap: 12px;
      border-bottom: 1px solid var(--border-soft);
      padding-bottom: 10px;
      margin-bottom: 12px;
    }}
    .dimension-head span {{
      color: var(--brand);
      font-size: 28px;
      line-height: 1;
    }}
    .table-wrap {{
      width: 100%;
      max-width: 100%;
      overflow-x: auto;
      margin: 14px 0 22px;
      border: 1px solid var(--border-soft);
      border-radius: 0;
    }}
    table {{
      width: 100%;
      border-collapse: collapse;
      min-width: 720px;
      font-size: 14px;
      line-height: 1.45;
    }}
    table.compact {{ font-size: 13px; }}
    th {{
      text-align: left;
      color: var(--dark-warm);
      font-weight: 500;
      background: #fbfbfa;
      border-bottom: 1px solid var(--border);
      padding: 10px 12px;
    }}
    td {{
      vertical-align: top;
      border-bottom: 1px solid var(--border-soft);
      padding: 9px 12px;
      overflow-wrap: anywhere;
    }}
    tr:last-child td {{ border-bottom: 0; }}
    ul {{ margin: 8px 0 0 20px; padding: 0; }}
    li {{ margin: 4px 0; }}
    p, li {{
      text-wrap: pretty;
    }}
    .fact-box p, .dimension-card p, .note, .chart-insight, .chart-recommendation {{
      text-align: justify;
      text-justify: inter-ideograph;
    }}
    .muted {{ color: var(--stone); }}
    .note {{
      margin-top: 14px;
      background: #fbfbfa;
      overflow-wrap: anywhere;
    }}
    .two-col {{
      display: grid;
      grid-template-columns: minmax(0, 1fr) minmax(0, 1fr);
      gap: 18px;
    }}
    .product-detail-grid {{
      display: grid;
      grid-template-columns: minmax(0, 1fr) minmax(0, 1fr);
      gap: 18px;
      margin-top: 22px;
      align-items: start;
    }}
    .product-detail-grid h3 {{
      margin-bottom: 10px;
    }}
    .chart-grid {{
      display: grid;
      grid-template-columns: repeat(2, minmax(0, 1fr));
      gap: 24px;
      align-items: stretch;
    }}
    .chart-module {{
      border: 1px solid var(--border);
      background: var(--paper);
      border-radius: 0;
      padding: 20px;
      break-inside: avoid;
      display: flex;
      flex-direction: column;
      height: 100%;
      min-width: 0;
    }}
    .chart-module-head {{
      display: grid;
      grid-template-columns: minmax(0, 1fr) minmax(150px, 220px);
      gap: 14px;
      border-bottom: 1px solid var(--border-soft);
      padding-bottom: 12px;
      margin-bottom: 14px;
    }}
    .chart-kicker {{
      display: inline-block;
      color: var(--brand);
      font-size: 12px;
      line-height: 1;
      margin-bottom: 7px;
      text-transform: uppercase;
      letter-spacing: 0.04em;
    }}
    .chart-confidence {{
      max-width: 220px;
      color: var(--stone);
      font-size: 12px;
      line-height: 1.35;
      text-align: right;
      overflow-wrap: anywhere;
    }}
    .chart-svg-wrap {{
      width: 100%;
      overflow: visible;
      margin: 10px 0 16px;
    }}
    .chart-svg-wrap svg {{
      display: block;
      width: 100%;
      min-width: 0;
      height: auto;
      background: #ffffff;
    }}
    .chart-insight, .chart-recommendation {{
      margin: 8px 0 0;
      color: var(--dark-warm);
      font-size: 14px;
      line-height: 1.62;
      overflow-wrap: anywhere;
    }}
    .chart-insight {{
      margin-top: auto;
    }}
    a {{ color: var(--brand); overflow-wrap: anywhere; }}
    @page {{
      size: A4;
      margin: 18mm 20mm 20mm 20mm;
      background: #ffffff;
    }}
    @media print {{
      body {{ padding-top: 0; font-size: 10pt; line-height: 1.52; }}
      .top-nav {{ display: none; }}
      main {{ max-width: none; padding: 0; }}
      section, .score-card, .fact-box, .risk-item {{ break-inside: avoid; }}
      h1 {{ font-size: 28pt; line-height: 1.12; }}
      h2 {{ font-size: 16pt; }}
      h3 {{ font-size: 12pt; }}
      .lede {{ max-width: none; font-size: 11pt; line-height: 1.58; }}
      .dimension-grid, .meta-grid, .score-grid, .fact-grid, .two-col, .product-detail-grid {{
        display: block;
      }}
      .chart-grid {{
        display: block;
      }}
      .chart-module {{
        margin-bottom: 13pt;
        padding: 12pt;
        border-color: #d8d5c9;
      }}
      .chart-module-head {{
        display: block;
      }}
      .chart-confidence {{
        max-width: none;
        text-align: left;
      }}
      .chart-svg-wrap {{
        overflow: visible;
      }}
      .chart-svg-wrap svg {{
        min-width: 0;
      }}
      .dimension-card, .score-card, .meta-item, .fact-box {{
        margin-bottom: 10pt;
      }}
      table {{ min-width: 0; font-size: 8.5pt; }}
      .table-wrap {{ overflow: visible; border: 0; }}
    }}
    @media (max-width: 820px) {{
      .nav-inner {{ padding: 10px 18px; display: block; }}
      .nav-title {{ max-width: none; margin-bottom: 8px; }}
      main {{ padding: 32px 18px 56px; }}
      .meta-grid, .score-grid, .fact-grid, .dimension-grid, .two-col, .chart-grid, .product-detail-grid {{
        grid-template-columns: minmax(0, 1fr);
      }}
      h1 {{ font-size: 34px; }}
      h2 {{ font-size: 23px; }}
      .chart-module-head {{
        display: block;
      }}
      .chart-confidence {{
        max-width: none;
        text-align: left;
        margin-top: 4px;
      }}
      .chart-svg-wrap {{
        overflow-x: auto;
        overflow-y: hidden;
      }}
      .chart-svg-wrap svg {{
        min-width: 620px;
      }}
      .chart-score_gauge .chart-svg-wrap svg {{
        min-width: 0;
      }}
      .chart-score_gauge .chart-svg-wrap {{
        overflow: visible;
      }}
      .chart-bar .chart-svg-wrap svg {{
        min-width: 0;
      }}
      .chart-bar .chart-svg-wrap {{
        overflow: visible;
      }}
      .chart-matrix .chart-svg-wrap svg {{
        min-width: 0;
      }}
      .chart-matrix .chart-svg-wrap {{
        overflow: visible;
      }}
    }}
  </style>
</head>
<body>
  <nav class="top-nav">
    <div class="nav-inner">
      <div class="nav-title">{h(meta.get('product_name') or title)}</div>
      <div class="nav-links">{nav_links}</div>
    </div>
  </nav>
  <main>
    <header class="report-header">
      <div class="eyebrow">需求三角评估报告</div>
      <h1>{h(title)}</h1>
      <p class="lede">{h(summary.get('one_sentence'))}</p>
      <div class="meta-grid">
        <div class="meta-item"><span>产品</span>{h(meta.get('product_name'))}</div>
        <div class="meta-item"><span>日期</span>{h(meta.get('generated_at'))}</div>
        <div class="meta-item"><span>目标市场</span>{h(meta.get('target_market', '未指定'))}</div>
        <div class="meta-item"><span>建议决策</span><span class="decision">{h(decision_label(summary.get('decision', '')))}</span></div>
      </div>
      <div class="score-grid">
        <div class="score-card total-score band-{band}"><span>总分</span><strong>{score(summary.get('total_score'))}</strong></div>
        <div class="score-card"><span>缺乏感</span><strong>{score(summary.get('lack_score'))}</strong></div>
        <div class="score-card"><span>目标物</span><strong>{score(summary.get('target_object_score'))}</strong></div>
        <div class="score-card"><span>消费者能力</span><strong>{score(summary.get('consumer_ability_score'))}</strong></div>
      </div>
    </header>

    <section id="summary">
      <h2>执行摘要</h2>
      <div class="two-col">
        <div class="fact-box"><h3>最大机会</h3><p>{h(summary.get('biggest_opportunity'))}</p></div>
        <div class="fact-box"><h3>最大风险</h3><p>{h(summary.get('biggest_risk'))}</p></div>
      </div>
      <div class="note"><strong>最短板：</strong>{h(weakest_dimension(report))}<br><strong>下一步：</strong>{h(final_plan.get('strategy') or summary.get('biggest_risk'))}</div>
    </section>

    <section id="visuals">
      <h2>可视化诊断</h2>
      <p class="lede">这些图表把需求强弱、短板、证据、风险和行动优先级压缩成可扫描的诊断模块。每个模块都包含解读和建议，避免只有图表没有判断。</p>
      <div class="chart-grid">
        {chart_html}
      </div>
    </section>

    <section id="product">
      <h2>产品概览</h2>
      <div class="fact-grid">
        <div class="fact-box"><h3>产品定义</h3><p>{h(canvas.get('definition'))}</p></div>
        <div class="fact-box"><h3>价值主张</h3><p>{h(canvas.get('value_proposition'))}</p></div>
        <div class="fact-box"><h3>商业模式</h3><p>{h(canvas.get('business_model'))}</p></div>
        <div class="fact-box"><h3>假设</h3>{html_list(as_list(canvas.get('assumptions')))}</div>
      </div>
      <div class="product-detail-grid">
        <div class="fact-box"><h3>核心功能</h3>{html_list(as_list(canvas.get('features')))}</div>
        <div class="fact-box"><h3>定价信息</h3>{html_list(as_list(canvas.get('pricing')))}</div>
      </div>
    </section>

    <section id="evidence">
      <h2>证据质量与来源</h2>
      <p>{h(meta.get('source_boundary', '本报告区分事实、假设、证据和建议。'))}</p>
      {html_table(['来源', '等级', '类型', '标题', '链接', '日期'], evidence_rows(report), compact=True)}
    </section>

    <section id="users">
      <h2>目标用户与 JTBD</h2>
      {html_table(['分群', '场景', 'JTBD', '当前替代', '采用阻碍'], segment_rows(report))}
    </section>

    <section id="competitors">
      <h2>竞品与替代方案</h2>
      {html_table(['名称', '类型', '定位', '优势', '弱点', '来源'], competitor_rows(report))}
    </section>

    <section id="triangle">
      <h2>需求三角分析</h2>
      <figure class="diagram-block triangle-figure">
        {render_triangle_svg(report)}
        <figcaption class="figure-caption">图 2：需求三角模型。缺乏感、目标物和消费者能力共同决定需求成立概率。</figcaption>
      </figure>
      <div class="dimension-grid">{''.join(dimension_cards)}</div>
    </section>

    <section id="scores">
      <h2>评分与解释</h2>
      {html_table(['维度', '分数'], score_rows(report), compact=True)}
    </section>

    <section id="recommendations">
      <h2>建议与实验</h2>
      {html_table(['优先级', '领域', '建议', '理由', '预期影响', '成本'], rec_rows)}
      <h3>验证实验</h3>
      {html_table(['假设', '分群', '方法', '指标', '阈值', '决策规则'], experiment_rows(report))}
    </section>

    <section id="forecast">
      <h2>预测情景</h2>
      <p>{h(forecast.get('recheck_trigger', '预测只作为情景推演，必须定期复盘。'))}</p>
      {html_table(['情景', '预测分数', '采用可能性', '关键假设'], [
        [item.get('name', ''), score(item.get('score_after')), item.get('adoption_likelihood', ''), '；'.join(as_list(item.get('assumptions')))]
        for item in as_list(forecast.get('scenarios')) if isinstance(item, dict)
      ])}
    </section>

    <section id="risks">
      <h2>风险与伦理</h2>
      {html_table(['严重度', '类型', '风险', '缓释', '来源'], risk_rows(report))}
    </section>

    <section id="final">
      <h2>最终方案</h2>
      <p class="lede">{h(final_plan.get('final_judgment', '未提供最终判断'))}</p>
      <div class="note"><strong>总体策略</strong><br>{h(final_plan.get('strategy', '未提供'))}</div>
      <div class="three-phase">
        {html_table(['阶段', '动作'], [
          ['未来 30 天', '；'.join(as_list(final_plan.get('next_30_days')))],
          ['未来 60 天', '；'.join(as_list(final_plan.get('next_60_days')))],
          ['未来 90 天', '；'.join(as_list(final_plan.get('next_90_days')))],
          ['决策规则', '；'.join(as_list(final_plan.get('decision_rules')))],
        ])}
      </div>
    </section>

    <section id="method">
      <h2>方法与模型</h2>
      <figure class="diagram-block">
        {render_process_svg()}
        <figcaption class="figure-caption">图：需求评估分析流程。检索只接受可追溯来源，不确定信息进入假设区。</figcaption>
      </figure>
      <p>评分采用需求三角的几何短板逻辑：缺乏感、目标物和消费者能力必须同时成立，单一高分不能完全补偿另一条边塌陷。</p>
    </section>

    <section id="appendix">
      <h2>附录</h2>
      <h3>未解问题</h3>
      {html_list(as_list(report.get('open_questions')))}
    </section>
  </main>
</body>
</html>
"""
    return html_doc


def set_docx_cell_shading(cell: Any, fill: str) -> None:
    try:
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn

        tc_pr = cell._tc.get_or_add_tcPr()
        shading = OxmlElement("w:shd")
        shading.set(qn("w:fill"), fill)
        tc_pr.append(shading)
    except Exception:
        return


def style_docx_table(table: Any) -> None:
    try:
        from docx.shared import Pt

        for row_index, row in enumerate(table.rows):
            for cell in row.cells:
                if row_index == 0:
                    set_docx_cell_shading(cell, "F7F7F4")
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.space_after = Pt(3)
                    paragraph.paragraph_format.line_spacing = 1.25
                    for run in paragraph.runs:
                        run.font.name = "Songti SC"
                        run.font.size = Pt(8.8 if len(table.columns) >= 5 else 9.5)
                        if row_index == 0:
                            run.bold = True
    except Exception:
        return


def add_docx_table(document: Any, headers: Sequence[str], rows: Sequence[Sequence[Any]]) -> None:
    if not rows:
        document.add_paragraph("未提供")
        return
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = True
    header_cells = table.rows[0].cells
    for index, header in enumerate(headers):
        header_cells[index].text = text(header)
    for row in rows:
        cells = table.add_row().cells
        for index, cell in enumerate(row):
            cells[index].text = text(cell)
    style_docx_table(table)


def add_docx_bullets(document: Any, items: Sequence[Any]) -> None:
    values = [text(item).strip() for item in items if text(item).strip()]
    if not values:
        document.add_paragraph("未提供")
        return
    for item in values:
        document.add_paragraph(item, style="List Bullet")


def add_docx_chart_module(document: Any, module: Dict[str, Any]) -> None:
    document.add_heading(text(module.get("title")), 2)
    inserted_image = False
    try:
        import cairosvg
        from docx.shared import Inches

        png_stream = io.BytesIO()
        cairosvg.svg2png(
            bytestring=render_chart_svg(module).encode("utf-8"),
            write_to=png_stream,
            output_width=900,
        )
        png_stream.seek(0)
        document.add_picture(png_stream, width=Inches(6.4))
        inserted_image = True
    except Exception:
        inserted_image = False
    if not inserted_image:
        document.add_paragraph("图表等价数据：")
        add_docx_table(document, ["分组", "指标/情景", "数值/X", "Y/说明"], chart_data_rows(module))
    document.add_paragraph(f"解读：{text(module.get('insight'))}")
    document.add_paragraph(f"建议：{text(module.get('recommendation'))}")
    document.add_paragraph(chart_meta(module))


def xml_text(value: Any) -> str:
    return html.escape(text(value), quote=False)


def ooxml_paragraph(value: Any, style: str | None = None) -> str:
    ppr_parts = []
    if style:
        ppr_parts.append(f'<w:pStyle w:val="{style}"/>')
    ppr_parts.append('<w:spacing w:after="120" w:line="330" w:lineRule="auto"/>')
    style_xml = f"<w:pPr>{''.join(ppr_parts)}</w:pPr>"
    return f"<w:p>{style_xml}<w:r><w:t>{xml_text(value)}</w:t></w:r></w:p>"


def ooxml_bullets(items: Sequence[Any]) -> str:
    values = [text(item).strip() for item in items if text(item).strip()]
    if not values:
        return ooxml_paragraph("未提供")
    return "".join(ooxml_paragraph(f"- {item}") for item in values)


def ooxml_table(headers: Sequence[str], rows: Sequence[Sequence[Any]]) -> str:
    if not rows:
        return ooxml_paragraph("未提供")
    table_rows = []
    all_rows = [headers] + [list(row) for row in rows]
    col_width = max(1200, int(9500 / max(1, len(headers))))
    for row_index, row in enumerate(all_rows):
        cells = []
        for cell in row:
            shading = '<w:shd w:fill="F7F7F4"/>' if row_index == 0 else ""
            cells.append(
                f"<w:tc><w:tcPr><w:tcW w:w=\"{col_width}\" w:type=\"dxa\"/>"
                f"{shading}<w:tcMar><w:top w:w=\"80\" w:type=\"dxa\"/>"
                "<w:left w:w=\"90\" w:type=\"dxa\"/><w:bottom w:w=\"80\" w:type=\"dxa\"/>"
                "<w:right w:w=\"90\" w:type=\"dxa\"/></w:tcMar></w:tcPr>"
                f"{ooxml_paragraph(cell)}</w:tc>"
            )
        table_rows.append(f"<w:tr>{''.join(cells)}</w:tr>")
    return (
        "<w:tbl><w:tblPr><w:tblStyle w:val=\"TableGrid\"/>"
        "<w:tblW w:w=\"9500\" w:type=\"dxa\"/><w:tblCellMar>"
        "<w:top w:w=\"80\" w:type=\"dxa\"/><w:left w:w=\"90\" w:type=\"dxa\"/>"
        "<w:bottom w:w=\"80\" w:type=\"dxa\"/><w:right w:w=\"90\" w:type=\"dxa\"/>"
        "</w:tblCellMar></w:tblPr>"
        f"{''.join(table_rows)}</w:tbl>"
    )


def fallback_docx_parts(body_xml: str) -> Dict[str, str]:
    document_xml = f"""<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"
  xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">
  <w:body>
    {body_xml}
    <w:sectPr>
      <w:pgSz w:w="11906" w:h="16838"/>
      <w:pgMar w:top="1440" w:right="1440" w:bottom="1440" w:left="1440" w:header="720" w:footer="720" w:gutter="0"/>
    </w:sectPr>
  </w:body>
</w:document>"""
    styles_xml = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:styles xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
  <w:style w:type="paragraph" w:default="1" w:styleId="Normal"><w:name w:val="Normal"/><w:pPr><w:spacing w:after="120" w:line="330" w:lineRule="auto"/></w:pPr><w:rPr><w:sz w:val="21"/><w:rFonts w:ascii="Times New Roman" w:eastAsia="Songti SC" w:hAnsi="Times New Roman"/></w:rPr></w:style>
  <w:style w:type="paragraph" w:styleId="Title"><w:name w:val="Title"/><w:basedOn w:val="Normal"/><w:pPr><w:spacing w:after="260"/></w:pPr><w:rPr><w:b/><w:color w:val="141413"/><w:sz w:val="44"/></w:rPr></w:style>
  <w:style w:type="paragraph" w:styleId="Heading1"><w:name w:val="heading 1"/><w:basedOn w:val="Normal"/><w:pPr><w:spacing w:before="300" w:after="160"/></w:pPr><w:rPr><w:b/><w:color w:val="1B365D"/><w:sz w:val="30"/></w:rPr></w:style>
  <w:style w:type="paragraph" w:styleId="Heading2"><w:name w:val="heading 2"/><w:basedOn w:val="Normal"/><w:pPr><w:spacing w:before="180" w:after="120"/></w:pPr><w:rPr><w:b/><w:color w:val="141413"/><w:sz w:val="26"/></w:rPr></w:style>
  <w:style w:type="paragraph" w:styleId="Heading3"><w:name w:val="heading 3"/><w:basedOn w:val="Normal"/><w:pPr><w:spacing w:before="120" w:after="80"/></w:pPr><w:rPr><w:b/><w:color w:val="3D3D3A"/><w:sz w:val="23"/></w:rPr></w:style>
  <w:style w:type="table" w:styleId="TableGrid"><w:name w:val="Table Grid"/><w:tblPr><w:tblBorders><w:top w:val="single" w:sz="6" w:color="E8E6DC"/><w:left w:val="single" w:sz="6" w:color="E8E6DC"/><w:bottom w:val="single" w:sz="6" w:color="E8E6DC"/><w:right w:val="single" w:sz="6" w:color="E8E6DC"/><w:insideH w:val="single" w:sz="4" w:color="EFEEE8"/><w:insideV w:val="single" w:sz="4" w:color="EFEEE8"/></w:tblBorders></w:tblPr></w:style>
</w:styles>"""
    return {
        "[Content_Types].xml": """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">
  <Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>
  <Default Extension="xml" ContentType="application/xml"/>
  <Override PartName="/word/document.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"/>
  <Override PartName="/word/styles.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.styles+xml"/>
</Types>""",
        "_rels/.rels": """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
  <Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" Target="word/document.xml"/>
</Relationships>""",
        "word/_rels/document.xml.rels": """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships"/>""",
        "word/document.xml": document_xml,
        "word/styles.xml": styles_xml,
    }


def render_docx_fallback(report: Dict[str, Any], output_path: Path) -> None:
    meta = report.get("meta", {})
    summary = report.get("executive_summary", {})
    canvas = report.get("product_canvas", {})
    triangle = report.get("triangle_analysis", {})
    forecast = report.get("forecast", {})
    final_plan = report.get("final_plan", {})
    body: List[str] = []

    body.append(ooxml_paragraph(meta.get("title") or meta.get("product_name") or "需求评估报告", "Title"))
    body.append(ooxml_paragraph(summary.get("one_sentence")))
    body.append(ooxml_table(["项目", "结果"], [
        ["产品", meta.get("product_name", "")],
        ["日期", meta.get("generated_at", "")],
        ["目标市场", meta.get("target_market", "")],
        ["建议决策", decision_label(summary.get("decision", ""))],
        ["总分", score(summary.get("total_score"))],
        ["缺乏感", score(summary.get("lack_score"))],
        ["目标物", score(summary.get("target_object_score"))],
        ["消费者能力", score(summary.get("consumer_ability_score"))],
        ["证据信心", summary.get("evidence_confidence", "")],
    ]))

    body.append(ooxml_paragraph("可视化诊断", "Heading1"))
    for module in chart_modules(report):
        body.append(ooxml_paragraph(module.get("title"), "Heading2"))
        body.append(ooxml_table(["分组", "指标/情景", "数值/X", "Y/说明"], chart_data_rows(module)))
        body.append(ooxml_paragraph(f"解读：{text(module.get('insight'))}"))
        body.append(ooxml_paragraph(f"建议：{text(module.get('recommendation'))}"))
        body.append(ooxml_paragraph(chart_meta(module)))

    body.append(ooxml_paragraph("产品概览", "Heading1"))
    body.append(ooxml_paragraph(f"产品定义：{text(canvas.get('definition'))}"))
    body.append(ooxml_paragraph(f"价值主张：{text(canvas.get('value_proposition'))}"))
    body.append(ooxml_paragraph(f"商业模式：{text(canvas.get('business_model'))}"))
    body.append(ooxml_paragraph("核心功能", "Heading2"))
    body.append(ooxml_bullets(as_list(canvas.get("features"))))
    body.append(ooxml_paragraph("定价信息", "Heading2"))
    body.append(ooxml_bullets(as_list(canvas.get("pricing"))))
    body.append(ooxml_paragraph("假设", "Heading2"))
    body.append(ooxml_bullets(as_list(canvas.get("assumptions"))))

    body.append(ooxml_paragraph("研究方法与来源", "Heading1"))
    body.append(ooxml_paragraph(meta.get("source_boundary", "本报告区分事实、假设、证据和建议。")))
    body.append(ooxml_paragraph("图 1：输入 -> 解析 -> 检索 -> 分析 -> 评分 -> 输出。"))
    body.append(ooxml_table(["来源", "等级", "类型", "标题", "链接", "日期"], evidence_rows(report)))

    body.append(ooxml_paragraph("目标用户与 JTBD", "Heading1"))
    body.append(ooxml_table(["分群", "场景", "JTBD", "当前替代", "采用阻碍"], segment_rows(report)))

    body.append(ooxml_paragraph("竞品与替代方案", "Heading1"))
    body.append(ooxml_table(["名称", "类型", "定位", "优势", "弱点", "来源"], competitor_rows(report)))

    body.append(ooxml_paragraph("需求三角分析", "Heading1"))
    body.append(ooxml_paragraph("图 2：缺乏感、目标物和消费者能力共同决定需求成立概率；任一维度明显缺失，总分都会下降。"))
    for key, label in [("lack", "缺乏感"), ("target_object", "目标物"), ("consumer_ability", "消费者能力")]:
        dim = triangle.get(key, {})
        body.append(ooxml_paragraph(f"{label}：{score(dim.get('score'))}", "Heading2"))
        body.append(ooxml_paragraph(f"推理：{text(dim.get('reasoning'))}"))
        body.append(ooxml_paragraph("支持证据", "Heading3"))
        body.append(ooxml_bullets(as_list(dim.get("evidence"))))
        body.append(ooxml_paragraph("反证或缺口", "Heading3"))
        body.append(ooxml_bullets(as_list(dim.get("counter_evidence"))))
        body.append(ooxml_paragraph(f"改进路径：{text(dim.get('improvement_path'))}"))

    body.append(ooxml_paragraph("评分与解释", "Heading1"))
    body.append(ooxml_table(["维度", "分数"], score_rows(report)))

    body.append(ooxml_paragraph("建议与实验", "Heading1"))
    rec_rows = [
        [
            item.get("priority", ""),
            item.get("area", ""),
            item.get("recommendation", ""),
            item.get("rationale", ""),
            item.get("expected_impact", ""),
            item.get("effort", ""),
        ]
        for item in as_list(report.get("recommendations"))
    ]
    body.append(ooxml_table(["优先级", "领域", "建议", "理由", "预期影响", "成本"], rec_rows))
    body.append(ooxml_paragraph("验证实验", "Heading2"))
    body.append(ooxml_table(["假设", "分群", "方法", "指标", "阈值", "决策规则"], experiment_rows(report)))

    body.append(ooxml_paragraph("风险与伦理", "Heading1"))
    body.append(ooxml_table(["严重度", "类型", "风险", "缓释", "来源"], risk_rows(report)))

    body.append(ooxml_paragraph("预测情景", "Heading1"))
    body.append(ooxml_paragraph(f"预测窗口：{text(forecast.get('horizon'))}；置信度：{text(forecast.get('confidence'))}"))
    body.append(ooxml_paragraph(f"复盘触发：{text(forecast.get('recheck_trigger'))}"))
    body.append(ooxml_table(
        ["情景", "预测分数", "采用可能性", "关键假设"],
        [
            [
                item.get("name", ""),
                score(item.get("score_after")),
                item.get("adoption_likelihood", ""),
                "；".join(as_list(item.get("assumptions"))),
            ]
            for item in as_list(forecast.get("scenarios"))
            if isinstance(item, dict)
        ],
    ))

    body.append(ooxml_paragraph("最终方案", "Heading1"))
    body.append(ooxml_paragraph(f"最终判断：{text(final_plan.get('final_judgment'))}"))
    body.append(ooxml_paragraph(f"总体策略：{text(final_plan.get('strategy'))}"))
    body.append(ooxml_table(["阶段", "动作"], [
        ["未来 30 天", "；".join(as_list(final_plan.get("next_30_days")))],
        ["未来 60 天", "；".join(as_list(final_plan.get("next_60_days")))],
        ["未来 90 天", "；".join(as_list(final_plan.get("next_90_days")))],
        ["决策规则", "；".join(as_list(final_plan.get("decision_rules")))],
    ]))

    body.append(ooxml_paragraph("附录", "Heading1"))
    body.append(ooxml_paragraph("未解问题", "Heading2"))
    body.append(ooxml_bullets(as_list(report.get("open_questions"))))

    parts = fallback_docx_parts("".join(body))
    with zipfile.ZipFile(output_path, "w", zipfile.ZIP_DEFLATED) as package:
        for name, content in parts.items():
            package.writestr(name, content)


def render_docx(report: Dict[str, Any], output_path: Path) -> None:
    try:
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
    except ImportError as exc:
        render_docx_fallback(report, output_path)
        return

    meta = report.get("meta", {})
    summary = report.get("executive_summary", {})
    canvas = report.get("product_canvas", {})
    triangle = report.get("triangle_analysis", {})
    forecast = report.get("forecast", {})
    final_plan = report.get("final_plan", {})
    document = Document()

    def set_style_font(style_name: str, size: float | None = None, color: str | None = None, bold: bool | None = None) -> None:
        style = styles[style_name]
        style.font.name = "Songti SC"
        if size is not None:
            style.font.size = Pt(size)
        if color is not None:
            style.font.color.rgb = RGBColor.from_string(color)
        if bold is not None:
            style.font.bold = bold
        r_pr = style._element.get_or_add_rPr()
        r_fonts = r_pr.rFonts
        if r_fonts is None:
            r_fonts = OxmlElement("w:rFonts")
            r_pr.append(r_fonts)
        r_fonts.set(qn("w:eastAsia"), "Songti SC")

    for section in document.sections:
        section.top_margin = Inches(0.72)
        section.bottom_margin = Inches(0.72)
        section.left_margin = Inches(0.78)
        section.right_margin = Inches(0.78)

    styles = document.styles
    set_style_font("Normal", 10.5, "141413", None)
    set_style_font("Title", 24, "141413", True)
    set_style_font("Heading 1", 16, "1B365D", True)
    set_style_font("Heading 2", 13, "141413", True)
    set_style_font("Heading 3", 11.5, "3D3D3A", True)

    normal = styles["Normal"].paragraph_format
    normal.line_spacing = 1.35
    normal.space_after = Pt(6)

    document.add_heading(text(meta.get("title") or meta.get("product_name") or "需求评估报告"), 0)
    lede = document.add_paragraph(text(summary.get("one_sentence")))
    lede.paragraph_format.line_spacing = 1.45
    lede.paragraph_format.space_after = Pt(12)
    add_docx_table(document, ["项目", "结果"], [
        ["产品", meta.get("product_name", "")],
        ["日期", meta.get("generated_at", "")],
        ["目标市场", meta.get("target_market", "")],
        ["建议决策", decision_label(summary.get("decision", ""))],
        ["总分", score(summary.get("total_score"))],
        ["缺乏感", score(summary.get("lack_score"))],
        ["目标物", score(summary.get("target_object_score"))],
        ["消费者能力", score(summary.get("consumer_ability_score"))],
        ["证据信心", summary.get("evidence_confidence", "")],
    ])

    document.add_heading("可视化诊断", 1)
    for module in chart_modules(report):
        add_docx_chart_module(document, module)

    document.add_heading("产品概览", 1)
    document.add_paragraph(f"产品定义：{text(canvas.get('definition'))}")
    document.add_paragraph(f"价值主张：{text(canvas.get('value_proposition'))}")
    document.add_paragraph(f"商业模式：{text(canvas.get('business_model'))}")
    document.add_heading("核心功能", 2)
    add_docx_bullets(document, as_list(canvas.get("features")))
    document.add_heading("定价信息", 2)
    add_docx_bullets(document, as_list(canvas.get("pricing")))
    document.add_heading("假设", 2)
    add_docx_bullets(document, as_list(canvas.get("assumptions")))

    document.add_heading("研究方法与来源", 1)
    document.add_paragraph(text(meta.get("source_boundary", "本报告区分事实、假设、证据和建议。")))
    document.add_paragraph("图 1：输入 -> 解析 -> 检索 -> 分析 -> 评分 -> 输出。检索只接受可追溯来源，不确定信息进入假设区。")
    add_docx_table(document, ["来源", "等级", "类型", "标题", "链接", "日期"], evidence_rows(report))

    document.add_heading("目标用户与 JTBD", 1)
    add_docx_table(document, ["分群", "场景", "JTBD", "当前替代", "采用阻碍"], segment_rows(report))

    document.add_heading("竞品与替代方案", 1)
    add_docx_table(document, ["名称", "类型", "定位", "优势", "弱点", "来源"], competitor_rows(report))

    document.add_heading("需求三角分析", 1)
    document.add_paragraph("图 2：需求三角模型。缺乏感、目标物和消费者能力共同决定需求成立概率；任一维度明显缺失，总分都会下降。")
    for key, label in [("lack", "缺乏感"), ("target_object", "目标物"), ("consumer_ability", "消费者能力")]:
        dim = triangle.get(key, {})
        document.add_heading(f"{label}：{score(dim.get('score'))}", 2)
        document.add_paragraph(f"推理：{text(dim.get('reasoning'))}")
        document.add_heading("支持证据", 3)
        add_docx_bullets(document, as_list(dim.get("evidence")))
        document.add_heading("反证或缺口", 3)
        add_docx_bullets(document, as_list(dim.get("counter_evidence")))
        document.add_paragraph(f"改进路径：{text(dim.get('improvement_path'))}")

    document.add_heading("评分与解释", 1)
    add_docx_table(document, ["维度", "分数"], score_rows(report))

    document.add_heading("建议与实验", 1)
    rec_rows = [
        [
            item.get("priority", ""),
            item.get("area", ""),
            item.get("recommendation", ""),
            item.get("rationale", ""),
            item.get("expected_impact", ""),
            item.get("effort", ""),
        ]
        for item in as_list(report.get("recommendations"))
    ]
    add_docx_table(document, ["优先级", "领域", "建议", "理由", "预期影响", "成本"], rec_rows)
    document.add_heading("验证实验", 2)
    add_docx_table(document, ["假设", "分群", "方法", "指标", "阈值", "决策规则"], experiment_rows(report))

    document.add_heading("风险与伦理", 1)
    add_docx_table(document, ["严重度", "类型", "风险", "缓释", "来源"], risk_rows(report))

    document.add_heading("预测情景", 1)
    document.add_paragraph(f"预测窗口：{text(forecast.get('horizon'))}；置信度：{text(forecast.get('confidence'))}")
    document.add_paragraph(f"复盘触发：{text(forecast.get('recheck_trigger'))}")
    add_docx_table(
        document,
        ["情景", "预测分数", "采用可能性", "关键假设"],
        [
            [
                item.get("name", ""),
                score(item.get("score_after")),
                item.get("adoption_likelihood", ""),
                "；".join(as_list(item.get("assumptions"))),
            ]
            for item in as_list(forecast.get("scenarios"))
            if isinstance(item, dict)
        ],
    )

    document.add_heading("最终方案", 1)
    document.add_paragraph(f"最终判断：{text(final_plan.get('final_judgment'))}")
    document.add_paragraph(f"总体策略：{text(final_plan.get('strategy'))}")
    add_docx_table(document, ["阶段", "动作"], [
        ["未来 30 天", "；".join(as_list(final_plan.get("next_30_days")))],
        ["未来 60 天", "；".join(as_list(final_plan.get("next_60_days")))],
        ["未来 90 天", "；".join(as_list(final_plan.get("next_90_days")))],
        ["决策规则", "；".join(as_list(final_plan.get("decision_rules")))],
    ])

    document.add_heading("附录", 1)
    document.add_heading("未解问题", 2)
    add_docx_bullets(document, as_list(report.get("open_questions")))

    document.save(output_path)


def render_pdf(html_doc: str, output_path: Path, base_url: Path) -> None:
    try:
        from weasyprint import HTML
    except ImportError as exc:
        raise RuntimeError("WeasyPrint is required for PDF output") from exc
    HTML(string=html_doc, base_url=str(base_url)).write_pdf(str(output_path))


def write_outputs(report: Dict[str, Any], outdir: Path, basename: str, formats: Sequence[str]) -> Dict[str, str]:
    outputs: Dict[str, str] = {}
    outdir.mkdir(parents=True, exist_ok=True)
    markdown_doc = render_markdown(report)
    html_doc = render_html(report)

    if "md" in formats:
        md_path = outdir / f"{basename}.md"
        md_path.write_text(markdown_doc, encoding="utf-8")
        outputs["markdown"] = str(md_path)
    if "html" in formats:
        html_path = outdir / f"{basename}.html"
        html_path.write_text(html_doc, encoding="utf-8")
        outputs["html"] = str(html_path)
    if "docx" in formats:
        docx_path = outdir / f"{basename}.docx"
        render_docx(report, docx_path)
        outputs["word"] = str(docx_path)
    if "pdf" in formats:
        pdf_path = outdir / f"{basename}.pdf"
        render_pdf(html_doc, pdf_path, outdir)
        outputs["pdf"] = str(pdf_path)
    return outputs


def parse_formats(value: str) -> List[str]:
    if value == "all":
        return ["md", "html", "docx", "pdf"]
    formats = [item.strip().lower() for item in value.split(",") if item.strip()]
    allowed = {"md", "html", "docx", "pdf"}
    unknown = sorted(set(formats) - allowed)
    if unknown:
        raise ValueError(f"unknown format(s): {', '.join(unknown)}")
    return formats


def main() -> None:
    parser = argparse.ArgumentParser(description="Render a Yao Demand Skill report to four output formats.")
    parser.add_argument("report_json", help="Path to canonical report JSON")
    parser.add_argument("--outdir", default=None, help="Output directory. Defaults to the report JSON directory.")
    parser.add_argument("--basename", default=None, help="Output basename. Defaults to slugified report title.")
    parser.add_argument("--formats", default="all", help="all, or comma list: md,html,docx,pdf")
    parser.add_argument("--skip-validate", action="store_true", help="Skip report validation before rendering")
    args = parser.parse_args()

    report_path = Path(args.report_json).resolve()
    report = json.loads(report_path.read_text(encoding="utf-8"))
    validation = validate_report(report)
    if not args.skip_validate and not validation["ok"]:
        print(json.dumps(validation, ensure_ascii=False, indent=2), file=sys.stderr)
        sys.exit(1)

    formats = parse_formats(args.formats)
    outdir = Path(args.outdir).resolve() if args.outdir else report_path.parent
    title = report.get("meta", {}).get("title") or report.get("meta", {}).get("product_name") or report_path.stem
    basename = args.basename or slugify(title)
    outputs = write_outputs(report, outdir, basename, formats)
    print(
        json.dumps(
            {
                "ok": True,
                "outputs": outputs,
                "warnings": validation.get("warnings", []),
            },
            ensure_ascii=False,
            indent=2,
        )
    )


if __name__ == "__main__":
    main()
